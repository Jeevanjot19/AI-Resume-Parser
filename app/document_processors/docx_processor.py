"""
DOCX document processor.
"""

from pathlib import Path
from typing import Dict, Any
import docx
from loguru import logger

from app.document_processors.base_processor import BaseProcessor


class DOCXProcessor(BaseProcessor):
    """DOCX document processor."""
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Add core properties if available
            if core_properties.author:
                metadata["author"] = core_properties.author
            if core_properties.created:
                metadata["created"] = core_properties.created.isoformat()
            if core_properties.modified:
                metadata["modified"] = core_properties.modified.isoformat()
            if core_properties.title:
                metadata["title"] = core_properties.title
            if core_properties.subject:
                metadata["subject"] = core_properties.subject
            
            return metadata
        except Exception as e:
            logger.error(f"DOCX metadata extraction error: {e}")
            return {"format": "docx"}
    
    def validate(self, file_path: Path) -> bool:
        """Validate DOCX file."""
        if not file_path.exists() or not file_path.is_file():
            return False
        
        if file_path.suffix.lower() not in ['.docx', '.doc']:
            return False
        
        try:
            docx.Document(file_path)
            return True
        except Exception:
            return False
